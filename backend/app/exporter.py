"""
Builds a clean, single-column, ATS-safe .docx from plain resume text.

This deliberately avoids tables, columns, text boxes, and graphics — the exact
things that break parsing on Workday/Taleo/iCIMS (see app/prompts.py). It uses
simple heuristics to detect section headers and bullet lines so the exported
file reads like a real resume, not a wall of plain text.
"""
import io
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT_NAME = "Calibri"
BODY_SIZE = Pt(10.5)
HEADER_SIZE = Pt(12)
NAME_SIZE = Pt(18)

# All-caps short lines (e.g. "EXPERIENCE", "EDUCATION") are treated as section headers.
_HEADER_RE = re.compile(r"^[A-Z][A-Z0-9 &/\-]{2,40}$")
_BULLET_RE = re.compile(r"^\s*[•\-\*\u2022]\s*(.*)")


def _set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(4)

    section = document.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)


def build_resume_docx(resume_text: str) -> bytes:
    document = Document()
    _set_default_style(document)

    lines = [line.rstrip() for line in resume_text.splitlines()]

    first_content_line_used = False

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            continue  # collapse blank lines — spacing is handled by paragraph space_after

        # First non-empty line is treated as the candidate's name — bigger, bold, centered.
        if not first_content_line_used:
            p = document.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = NAME_SIZE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_content_line_used = True
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            p = document.add_paragraph(style="List Bullet")
            p.add_run(bullet_match.group(1))
            continue

        if _HEADER_RE.match(line) and len(line.split()) <= 5:
            p = document.add_paragraph()
            run = p.add_run(line.upper())
            run.bold = True
            run.font.size = HEADER_SIZE
            p.paragraph_format.space_before = Pt(10)
            # simple bottom border effect via underline, since real borders need XML
            run.underline = True
            continue

        # default: plain paragraph
        document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
