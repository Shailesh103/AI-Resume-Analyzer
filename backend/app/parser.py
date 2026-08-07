"""
Extracts raw text from an uploaded resume file.
Supports PDF, DOCX, and plain text.
"""
import io
import pdfplumber
from docx import Document
from fastapi import HTTPException


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Route to the right extractor based on file extension."""
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    elif lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload a PDF, DOCX, or TXT resume.",
        )


def _extract_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not read PDF: {e}")

    text = "\n".join(text_chunks).strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="No extractable text found in this PDF. It may be a scanned "
            "image — try exporting your resume as a text-based PDF.",
        )
    return text


def _extract_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not read DOCX: {e}")

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables (many resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(status_code=422, detail="No text found in this DOCX file.")
    return text
